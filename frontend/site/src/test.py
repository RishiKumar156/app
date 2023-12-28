import json
from docx import Document

def json_to_docx(json_data, output_docx):
    document = Document()

    for key, value in json_data.items():
        document.add_heading(key, level=1)
        document.add_paragraph(str(value))

    document.save(output_docx)

# Example usage
json_data = '{"Title": "Document Title", "Content": "This is some content."}'
data = json.loads(json_data)
json_to_docx(data, 'output.docx')

def create_promissory_note(data, output_docx):
    document = Document()

    for key, value in data.items():
        document.add_paragraph(f"{key.upper()}:\n{value}\n")

    document.save(output_docx)

# Example usage
promissory_data = {
    "4. PREPAYMENT": "The Borrower has the right to pay back the loan in-full or make additional payments at any time without penalty.",
    "5. REMEDIES": "No delay or omission on part of the holder of this Note in exercising any right hereunder shall operate as a waiver of any such right or of any other right of such holder, nor shall any delay, omission or waiver on any one occasion be deemed a bar to or waiver of the same or any other right on any future occasion. The rights and remedies of the Lender shall be cumulative and may be pursued singly, successively, or together, in the sole discretion of the Lender.",
    "6. EVENTS OF ACCELERATION": "The occurrence of any of the following shall constitute an 'Event of Acceleration' by the Lender under this Note:\n\nBorrower’s failure to pay any part of the principal or interest as and when due under this Note; or\nBorrower’s becoming insolvent or not paying its debts as they become due.",
    # Add other sections as needed
}

create_promissory_note(promissory_data, 'promissory_note.docx')

def create_loan_agreement(loan_data, output_docx):
    document = Document()

    # Add loan agreement content to the document
    for section in loan_data:
        document.add_heading(section['heading'], level=1)
        
        for item in section['content']:
            if isinstance(item, dict):
                # If it's a dictionary, add a bullet list
                for key, value in item.items():
                    document.add_paragraph(f"{key}: {value}", style='ListBullet')
            else:
                # If it's a string, add it as a paragraph
                document.add_paragraph(item)

    # Save the document
    document.save(output_docx)

# Define the loan agreement data
loan_agreement_data = [
    {
        'heading': 'Loan Amount',
        'content': ['Ex: "TWO-THOUSAND" Dollars ($2,000.00)']
    },
    {
        'heading': 'Date',
        'content': ['\t\t\t\t']
    },
    # Add other sections as needed
    # ...
]

# Example usage
create_loan_agreement(loan_agreement_data, 'loan_agreement.docx')
